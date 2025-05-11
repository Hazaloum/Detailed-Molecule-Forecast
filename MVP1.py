import streamlit as st
import pandas as pd
import io
import base64
import os
import json
from docx import Document
from docx.shared import Pt
import contextlib
import sys

# --- Load Data ---
df = pd.read_csv("GoldMine Second Try.csv")
mapping_df = pd.read_csv("HGDR_with_forecast2.csv")

# Clean columns
df.columns = df.columns.str.replace("\n", " ", regex=False).str.strip()
mapping_df.columns = mapping_df.columns.str.replace("\n", " ", regex=False).str.strip()

# Clean numeric columns
for col in df.columns:
    if "Value" in col or "Units" in col:
        df[col] = pd.to_numeric(df[col].astype(str).str.replace(",", "").str.strip(), errors='coerce').fillna(0)

# --- External Data ---
external_data = {
    "CAGR": mapping_df.set_index("Molecule")["CAGR"].to_dict(),
    "PrivateShare": mapping_df.set_index("Molecule")["Private"].to_dict(),
    "Competitors": mapping_df.set_index("Molecule")["Comp"].to_dict()
}

molecule_list = sorted(df["Molecule"].dropna().unique().tolist())

# --- UI Layout ---
st.set_page_config(page_title="Pharma Analytics MVP", layout="wide")
tabs = st.tabs(["Overview & Products", "ATC4 Competition", "Shortlist Builder"])

# --- Tab 1: Overview + Products ---
with tabs[0]:
    st.header("🔍 Molecule Overview + Packs")
    selected_molecule = st.selectbox("Select Molecule:", molecule_list)

    if selected_molecule:
        mol_df = df[df["Molecule"] == selected_molecule]
        if not mol_df.empty:
            atc1 = mol_df["ATC1"].mode()[0]
            atc2 = mol_df["ATC2"].mode()[0]
            atc3 = mol_df["ATC3"].mode()[0]
            atc4 = mol_df["ATC4"].mode()[0]

            units_22, units_23, units_24 = mol_df["2022 Units"].sum(), mol_df["2023 Units"].sum(), mol_df["2024 Units"].sum()
            value_22, value_23, value_24 = mol_df["2022 LC Value"].sum(), mol_df["2023 LC Value"].sum(), mol_df["2024 LC Value"].sum()

            cagr = external_data["CAGR"].get(selected_molecule, 0)
            private = external_data["PrivateShare"].get(selected_molecule, 0)
            competitors = external_data["Competitors"].get(selected_molecule, 0)
            first_launch = int(mol_df["Launch Year"].min())

            atc4_df = df[df["ATC4"] == atc4]
            atc4_units_22 = atc4_df["2022 Units"].sum()
            atc4_units_23 = atc4_df["2023 Units"].sum()
            atc4_units_24 = atc4_df["2024 Units"].sum()
            atc4_value_22 = atc4_df["2022 LC Value"].sum()
            atc4_value_23 = atc4_df["2023 LC Value"].sum()
            atc4_value_24 = atc4_df["2024 LC Value"].sum()

            share_22 = units_22 / atc4_units_22 * 100
            share_23 = units_23 / atc4_units_23 * 100
            share_24 = units_24 / atc4_units_24 * 100

            def compute_cagr(start, end, years=2):
                try:
                    return ((end / start) ** (1 / years) - 1) * 100
                except ZeroDivisionError:
                    return 0

            units_cagr = compute_cagr(units_22, units_24)
            value_cagr = compute_cagr(value_22, value_24)

            st.markdown(f"""
            ### 🧪 Molecule: {selected_molecule}
            **🔬 Class**: {atc1} / {atc2} / {atc3} / {atc4}  
            **📦 Units 2022/23/24**: {int(units_22):,} / {int(units_23):,} / {int(units_24):,}  
            **📦 CAGR**: {cagr:.0f}%  
            **💸 Value 2024**: {int(value_24):,} AED  
            **🏷️ Private Share**: {private*100:.1f}%  
            **🤝 Competitors**: {competitors}  
            **📜 First Launch**: {first_launch}  
            **📊 ATC4 Unit Share**: {share_24:.2f}%  
            **📊 ATC4 Unit Share (2022/23/24)**: {share_22:.2f}% / {share_23:.2f}% / {share_24:.2f}%  
            **💸 Value 2022/23/24**: {int(value_22):,} / {int(value_23):,} / {int(value_24):,} AED  
            **🏛️ ATC4 Units 2022/23/24**: {int(atc4_units_22):,} / {int(atc4_units_23):,} / {int(atc4_units_24):,}  
            **🏛️ ATC4 Values 2022/23/24**: {int(atc4_value_22):,} / {int(atc4_value_23):,} / {int(atc4_value_24):,} AED  
            **📈 Units CAGR (22→24)**: {units_cagr:.2f}%  
            **📈 Value CAGR (22→24)**: {value_cagr:.2f}%  
            """)

            mol_units = mol_df["2024 Units"].sum()
            for product, prod_df in mol_df.groupby("Product"):
                product_units = prod_df["2024 Units"].sum()
                manufacturer = prod_df["Manufacturer"].mode()[0]
                combo_type = prod_df["Molecule Combination Type"].mode()[0]
                product_unit_share = product_units / mol_units * 100

                st.markdown(f"""
                **🧪 {product} — {manufacturer} — {combo_type} — {product_unit_share:.2f}% of molecule**  
                {"🧬 Mono-molecule Product" if len(prod_df["Molecule"].unique()) == 1 else "🔗 Shared Product"}
                """)

                for _, pack in prod_df.iterrows():
                    st.markdown(f" • {pack['Pack']} — {pack['Retail Price']:.2f} AED — {int(pack['2024 Units']):,} units")

# --- Tab 2: ATC4 Competition ---
with tabs[1]:
    st.header("🧭 ATC4 Competitor Map")
    mol_df = df[df["Molecule"] == selected_molecule]
    atc4 = mol_df["ATC4"].mode()[0]
    atc4_df = df[df["ATC4"] == atc4]

    grouped = (
        atc4_df.groupby("Molecule")
        .agg({"2024 Units": "sum", "2024 LC Value": "sum"})
        .sort_values(by="2024 Units", ascending=False)
        .reset_index()
    )
    total_units = atc4_df["2024 Units"].sum() or 1
    total_value = atc4_df["2024 LC Value"].sum() or 1
    grouped["% of ATC4 Units"] = (grouped["2024 Units"] / total_units * 100).round(2)
    grouped["% of ATC4 Value"] = (grouped["2024 LC Value"] / total_value * 100).round(2)
    grouped["Competitors"] = grouped["Molecule"].map(external_data["Competitors"]).fillna("N/A")
    st.dataframe(grouped)

# --- Tab 3: Shortlist Builder ---
with tabs[2]:
    st.header("🧩 Molecule Shortlist Builder")

    shortlist_file = "shortlist.txt"
    shortlist = json.load(open(shortlist_file)) if os.path.exists(shortlist_file) else []

    selected_add = st.selectbox("Search & Add Molecule", options=molecule_list)
    if st.button("➕ Add to shortlist"):
        if selected_add and selected_add not in shortlist:
            shortlist.append(selected_add)
            with open(shortlist_file, "w") as f:
                json.dump(shortlist, f)
            st.success(f"Added {selected_add} to shortlist.")

    if st.button("🗑️ Clear shortlist"):
        shortlist = []
        with open(shortlist_file, "w") as f:
            json.dump(shortlist, f)
        st.info("Shortlist cleared.")

    st.subheader("📋 Current Shortlist")
    st.write(shortlist)

    if st.button("📤 Export to Word"):
        doc = Document()
        doc.add_heading("📊 Molecule Portfolio Report", level=1)

        for mol in shortlist:
            mol_df = df[df["Molecule"] == mol]
            if mol_df.empty:
                continue

            atc1 = mol_df["ATC1"].mode()[0]
            atc2 = mol_df["ATC2"].mode()[0]
            atc3 = mol_df["ATC3"].mode()[0]
            atc4 = mol_df["ATC4"].mode()[0]
            units_22 = mol_df["2022 Units"].sum()
            units_23 = mol_df["2023 Units"].sum()
            units_24 = mol_df["2024 Units"].sum()
            value_22 = mol_df["2022 LC Value"].sum()
            value_23 = mol_df["2023 LC Value"].sum()
            value_24 = mol_df["2024 LC Value"].sum()

            cagr = external_data["CAGR"].get(mol, 0)
            private = external_data["PrivateShare"].get(mol, 0)
            competitors = external_data["Competitors"].get(mol, 0)
            first_launch = int(mol_df["Launch Year"].min())

            atc4_df = df[df["ATC4"] == atc4]
            atc4_units_22 = atc4_df["2022 Units"].sum()
            atc4_units_23 = atc4_df["2023 Units"].sum()
            atc4_units_24 = atc4_df["2024 Units"].sum()
            atc4_value_22 = atc4_df["2022 LC Value"].sum()
            atc4_value_23 = atc4_df["2023 LC Value"].sum()
            atc4_value_24 = atc4_df["2024 LC Value"].sum()

            share_22 = units_22 / atc4_units_22 * 100
            share_23 = units_23 / atc4_units_23 * 100
            share_24 = units_24 / atc4_units_24 * 100

            def compute_cagr(start, end, years=2):
                try:
                    return ((end / start) ** (1 / years) - 1) * 100
                except ZeroDivisionError:
                    return 0

            units_cagr = compute_cagr(units_22, units_24)
            value_cagr = compute_cagr(value_22, value_24)

            doc.add_heading(mol, level=2)
            p = doc.add_paragraph()
            p.add_run(f"Class: {atc1} / {atc2} / {atc3} / {atc4}\n")
            p.add_run(f"Units 22/23/24: {units_22:,} / {units_23:,} / {units_24:,}\n")
            p.add_run(f"CAGR: {cagr:.0f}%\n")
            p.add_run(f"Value 2024: {value_24:,} AED\n")
            p.add_run(f"Private Share: {private*100:.1f}%\n")
            p.add_run(f"Competitors: {competitors}\n")
            p.add_run(f"First Launch: {first_launch}\n")
            p.add_run(f"ATC4 Share 2022/23/24: {share_22:.2f}% / {share_23:.2f}% / {share_24:.2f}%\n")
            p.add_run(f"Value 22/23/24: {value_22:,} / {value_23:,} / {value_24:,} AED\n")
            p.add_run(f"ATC4 Units 22/23/24: {atc4_units_22:,} / {atc4_units_23:,} / {atc4_units_24:,}\n")
            p.add_run(f"ATC4 Values 22/23/24: {atc4_value_22:,} / {atc4_value_23:,} / {atc4_value_24:,} AED\n")
            p.add_run(f"Units CAGR (22→24): {units_cagr:.2f}%\n")
            p.add_run(f"Value CAGR (22→24): {value_cagr:.2f}%\n")

            mol_units = mol_df["2024 Units"].sum()
            for product, prod_df in mol_df.groupby("Product"):
                product_units = prod_df["2024 Units"].sum()
                manufacturer = prod_df["Manufacturer"].mode()[0]
                combo_type = prod_df["Molecule Combination Type"].mode()[0]
                product_unit_share = product_units / mol_units * 100

                p.add_run(f"\n🧪 {product} — {manufacturer} — {combo_type} — {product_unit_share:.2f}% of molecule\n")
                for _, pack in prod_df.iterrows():
                    p.add_run(f" • {pack['Pack']} — {pack['Retail Price']:.2f} AED — {int(pack['2024 Units']):,} units\n")

        filename = "Molecule_Portfolio_Report.docx"
        doc.save(filename)
        with open(filename, "rb") as f:
            st.download_button(label="📥 Download Portfolio Report", data=f, file_name=filename, mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
